import os
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

# Define the 5 distinct testing scenarios representing all edge cases
TEST_SUITES = {
    "test_1_prose": {
        "title": "Scenario 1: Standard Academic Prose and Deep Architecture",
        "content": [
            "Artificial Intelligence has transitioned from a theoretical academic discipline into the backbone of modern software engineering. In the early days, systems relied heavily on hardcoded rules and expert systems to simulate human intelligence. However, these systems were fragile and failed to scale when faced with real-world ambiguity.",
            "The major breakthrough arrived with the advent of deep learning and neural networks. By processing massive datasets, models learned to recognize patterns autonomously! Today, large language models can generate complex application code, analyze dense legal documents, and even assist doctors in diagnosing rare medical conditions? The field is moving incredibly fast."
        ],
        "type": "prose"
    },
    "test_2_config": {
        "title": "Scenario 2: Technical System Configuration Environment Setup",
        "content": [
            "Ensure your local deployment environment meets the following baseline parameters before running the RAG pipeline orchestrator node:",
            "SERVER_PORT=8080",
            "DB_TIMEOUT=3000",
            "LOG_LEVEL=DEBUG",
            "MAX_CONNECTIONS=100",
            "Note that changing these environment values without proper administrative tokens might isolate the microservice from the central secure gateway completely."
        ],
        "type": "config"
    },
    "test_3_bullets": {
        "title": "Scenario 3: Architectural Roadmap and Key Deliverables",
        "content": [
            "The infrastructure engineering team must prioritize the development and validation of these security and performance features:",
            "• Implement OAuth2 and OpenID Connect for secure machine-to-machine identity validation.",
            "• Deploy highly available Kubernetes pods across multiple distinct geographic cloud regions.",
            "• Integrate Model Context Protocol (MCP) servers to allow structural tool-calling capabilities.",
            "- Conduct extensive zero-trust network packet analysis under simulated vector database query loads.",
            "Failure to deliver these core components before the semester audit will result in a deployment hold."
        ],
        "type": "bullets"
    },
    "test_4_edge_cases": {
        "title": "Scenario 4: Stress Testing Text Extraction and Tokenizer Boundaries",
        "content": [
            "Dr. Hasson visited the U.S.A. at 5 A.M. to meet senior AI research scientists regarding multi-agent vector search logic.",
            "This sentence contains a massive cluster of consecutive spaces          that the clean function should resolve.",
            "Does the sentence splitter handle multiple exclamation marks properly??? Yes, it absolutely should!!!",
            "The following chunk is separated by multiple consecutive empty lines to stress test the parser newline handling logic."
        ],
        "type": "edge_cases"
    },
    "test_5_mixed_report": {
        "title": "Scenario 5: Combined Corporate Status Report Q2",
        "content": [
            "Executive Summary",
            "The engineering department completed the integration of the vector database abstraction layer ahead of schedule. Performance benchmarks indicate a thirty percent reduction in retrieval latency.",
            "Current System Vulnerabilities",
            "• Missing rate limiting infrastructure on the public LLM route endpoint.",
            "• Database connection pooling issues under spike workloads.",
            "Next Steps and Operational Milestones",
            "Refactor the routing layer and deploy to the staging cluster by Friday night."
        ],
        "type": "mixed"
    }
}

def generate_docx_suite():
    """Generates 5 distinct DOCX files based on the scenarios."""
    print("-> Generating 5 DOCX test files...")
    for filename, data in TEST_SUITES.items():
        doc = Document()
        doc.add_heading(data["title"], level=1)
        
        for p_text in data["content"]:
            if p_text.startswith("•") or p_text.startswith("-"):
                doc.add_paragraph(p_text[1:].strip(), style='ListBullet')
            else:
                doc.add_paragraph(p_text)
                
        full_path = f"{filename}.docx"
        doc.save(full_path)
        print(f"   Created: {full_path}")

def generate_pdf_suite():
    """Generates 5 matching PDF files based on the scenarios using ReportLab."""
    print("-> Generating 5 PDF test files...")
    styles = getSampleStyleSheet()
    
    # Custom styles to look like clean paragraphs
    title_style = ParagraphStyle('TitleStyle', parent=styles['Heading1'], spaceAfter=12)
    body_style = ParagraphStyle('BodyStyle', parent=styles['Normal'], spaceAfter=10, leading=14)
    
    for filename, data in TEST_SUITES.items():
        pdf_path = f"{filename}.pdf"
        doc = SimpleDocTemplate(pdf_path, pagesize=letter)
        story = []
        
        # Add Title
        story.append(Paragraph(data["title"], title_style))
        story.append(Spacer(1, 12))
        
        # Add paragraphs/bullets
        for p_text in data["content"]:
            # Reportlab handles bullets nicely if we format them with indentation
            if p_text.startswith("•") or p_text.startswith("-"):
                formatted_text = f"<b>&bull;</b> {p_text[1:].strip()}"
                bullet_style = ParagraphStyle('BulletStyle', parent=body_style, leftIndent=20)
                story.append(Paragraph(formatted_text, bullet_style))
            else:
                story.append(Paragraph(p_text, body_style))
                
        doc.build(story)
        print(f"   Created: {pdf_path}")

if __name__ == "__main__":
    print("============================================================")
    print("Starting Automated Generation of RAG Test Suite Matrix")
    print("============================================================")
    generate_docx_suite()
    print("-" * 60)
    generate_pdf_suite()
    print("============================================================")
    print("[SUCCESS] Created 10 matching test files (5 DOCX, 5 PDF)!")
    print("============================================================")